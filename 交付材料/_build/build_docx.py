from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "酒店管理系统_全套软件工程文档.docx"
AUTHOR = "XXX"
DATE = "2026-07-02"
TEST_RUN_ID = "019f21ac-5bd7-70f0-ada5-30be4ba097be"
DOC_FILES = [
    "README.md",
    "01-需求规格说明书.md",
    "02-概要设计说明书.md",
    "03-详细设计说明书.md",
    "04-数据库设计说明书.md",
    "05-接口设计说明书.md",
    "06-用户手册.md",
    "07-部署运行手册.md",
    "08-软件测试报告.md",
    "09-答辩材料清单.md",
    "10-演示视频脚本和录制清单.md",
    "11-格式排版查重自查表.md",
]


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = "F2F4F7"
BORDER = "B8BCC4"


def set_run_font(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "酒店管理系统软件工程文档"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in header.runs:
            set_run_font(run, size=9, color=MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = f"{AUTHOR} | {DATE} | 测试记录 {TEST_RUN_ID}"
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in footer.runs:
            set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("酒店管理系统")
    set_run_font(run, size=28, bold=True, color=RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("全套软件工程文档")
    set_run_font(run, size=16, color=MUTED)

    rows = [
        ("项目", "酒店管理系统"),
        ("负责人", AUTHOR),
        ("角色", "测试工程师 / 文档助理"),
        ("日期", DATE),
        ("测试记录编号", TEST_RUN_ID),
        ("说明", "Markdown 源文件与 Word 成品同步维护，PPT 初稿和演示脚本另行提供。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.text = value
            for run in para.runs:
                set_run_font(run, size=10.5, bold=(col_index == 0))

    doc.add_page_break()
    toc_title = doc.add_paragraph("文档目录", style="Heading 1")
    toc_title.paragraph_format.space_before = Pt(0)
    for idx, name in enumerate(DOC_FILES, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{idx}. {name}")
    doc.add_page_break()


def split_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    return table_lines, i


def parse_table(table_lines):
    rows = []
    for line in table_lines:
        parts = [part.strip() for part in line.strip("|").split("|")]
        if all(set(part.replace(" ", "")) <= {"-", ":"} for part in parts):
            continue
        rows.append(parts)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(max_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            cell.paragraphs[0].text = text
            for run in cell.paragraphs[0].runs:
                set_run_font(run, size=9.5, bold=(r_idx == 0))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(4)


def add_code_block(doc, block):
    text = "\n".join(block)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, font="Consolas")


inline_code = re.compile(r"`([^`]+)`")


def add_inline_text(paragraph, text, bold=False):
    pos = 0
    for match in inline_code.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=11, bold=bold)
        code = paragraph.add_run(match.group(1))
        set_run_font(code, size=10.5, bold=bold, font="Consolas")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=11, bold=bold)


def add_markdown_doc(doc, path: Path, first=False):
    if not first:
        doc.add_page_break()
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_block = []
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_block)
                code_block = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_block.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines, i = split_table(lines, i)
            add_table(doc, parse_table(table_lines))
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:], style="Heading 1")
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 2")
        elif stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 3")
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_text(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.space_after = Pt(3)
            add_inline_text(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_inline_text(p, stripped)
        i += 1


def main():
    doc = Document()
    configure_document(doc)
    doc.core_properties.author = AUTHOR
    doc.core_properties.title = "酒店管理系统全套软件工程文档"
    doc.core_properties.subject = "课程设计交付材料"
    add_cover(doc)
    for index, filename in enumerate(DOC_FILES):
        add_markdown_doc(doc, ROOT / filename, first=(index == 0))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
